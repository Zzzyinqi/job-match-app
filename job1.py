import streamlit as st
import os
import re
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from datetime import datetime
import difflib

# 创建上传目录
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 预定义的岗位数据库
JOB_DATABASE = {
    "技术岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "后端开发工程师",
        "学历要求": "本科及以上",
        "薪资范围": "15k-25k",
        "工作经验要求": "3年",
        "性别要求": "不限",
    },
    "法务岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "法务专员",
        "学历要求": "法学本科及以上",
        "薪资范围": "15k-20k",
        "工作经验要求": "3年",
        "性别要求": "不限",
    },
    "人事岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "人力资源专员",
        "学历要求": "大专及以上",
        "薪资范围": "8k-12k",
        "工作经验要求": "2年",
        "性别要求": "不限",
    },
    "行政岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "行政助理",
        "学历要求": "大专及以上",
        "薪资范围": "6k-9k",
        "工作经验要求": "1年",
        "性别要求": "不限",
    },
    "运营岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "内容运营专员",
        "学历要求": "本科及以上",
        "薪资范围": "10k-15k",
        "工作经验要求": "2年",
        "性别要求": "不限",
    },
    "后勤岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "后勤主管",
        "学历要求": "高中及以上",
        "薪资范围": "10k-14k",
        "工作经验要求": "5年",
        "性别要求": "不限",
    },
    "销售岗": {
        "企业名称": "星辰科技有限公司",
        "招聘岗位": "销售经理",
        "学历要求": "高中及以上",
        "薪资范围": "底薪8k+提成",
        "工作经验要求": "3年",
        "性别要求": "不限",
    }
}

def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本"""
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    """从DOCX文件中提取文本"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return "\n".join(full_text)

def extract_text_from_txt(file_path):
    """从TXT文件中提取文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def normalize_position(position):
    """标准化岗位名称"""
    if not position:
        return ""
    
    # 常见岗位名称映射
    position_mapping = {
        '前端': '前端开发', 'web前端': '前端开发', '前端工程师': '前端开发',
        '后端': '后端开发','后端': '后端开发工程师', 'java': 'Java开发', 'python': 'Python开发',
        '测试': '软件测试', 'qa': '软件测试', '测试工程师': '软件测试',
        '产品': '产品经理', '产品设计': '产品经理', 'pm': '产品经理',
        '运营': '运营专员', '新媒体': '新媒体运营', '内容运营': '内容运营',
        '销售': '销售代表',  '销售': '销售经理','业务员': '销售代表', 'bd': '商务拓展',
        '人事': '人力资源', 'hr': '人力资源', '招聘': '招聘专员',
        '财务': '财务会计', '会计': '财务会计', '出纳': '财务会计',
        '行政': '行政专员', '文员': '行政专员', '助理': '行政助理'
    }
    
    # 移除无关字符
    clean_position = re.sub(r'[、/（）()【】\[\]\s]', '', position).lower()
    
    # 应用映射
    for key, value in position_mapping.items():
        if key in clean_position:
            return value
    
    # 如果没有匹配的映射，返回原始值
    return position

def parse_document(text):
    """从求职者简历文本中提取结构化信息"""
    info = {
        '姓名': '',
        '年龄': '',
        '性别': '',
        '学历': '',
        '专业': '',
        '工作经验': '',
        '期望薪资': '',
        '求职岗位': '',
        '联系方式': ''
    }
    
    # 改进的姓名提取 - 匹配表格格式和无冒号格式
    name_match = re.search(r'(?:姓名|名字|个人姓名|候选人姓名)[\s:：]*([\u4e00-\u9fa5A-Za-z·]{2,4})', text)
    if not name_match:
        name_match = re.search(r'^[\s]*([\u4e00-\u9fa5]{2,4})[\s]*$', text, re.MULTILINE)
    if name_match:
        info['姓名'] = name_match.group(1)
    
    # 改进的年龄提取 - 匹配各种格式
    age_match = re.search(r'(?:年龄|岁数|出生年份)[\s:：]*(\d+)', text)
    if not age_match:
        age_match = re.search(r'(\d+)[\s]*(?:岁|years?|y/o)', text, re.IGNORECASE)
    if age_match:
        info['年龄'] = age_match.group(1)
    
    # 改进的性别提取 - 匹配各种格式
    gender_match = re.search(r'(?:性别)[\s:：]*([男女])', text)
    if not gender_match:
        gender_match = re.search(r'([男女])(?:\s*性)?', text)
    if gender_match:
        info['性别'] = gender_match.group(1)
    
    # 改进的学历提取 - 匹配各种格式
    education_match = re.search(r'(?:学历|教育背景|最高学历)[\s:：]*([\u4e00-\u9fa5]{2,4})', text)
    if not education_match:
        education_match = re.search(r'(本科|硕士|博士|大专|高中|中专|初中|小学)', text)
    if education_match:
        info['学历'] = education_match.group(1)
    
    # 改进的专业提取 - 匹配各种格式
    major_match = re.search(r'(?:专业|所学专业|主修专业)[\s:：]*([\u4e00-\u9fa5A-Za-z]{2,10})', text)
    if not major_match:
        major_match = re.search(r'专业[\s]*([\u4e00-\u9fa5A-Za-z]{2,10})', text)
    if major_match:
        info['专业'] = major_match.group(1)
    
    # 改进的工作经验提取 - 匹配各种格式
    exp_match = re.search(r'(?:工作经验|工作年限|从业时间)[\s:：]*(\d+)', text)
    if not exp_match:
        exp_match = re.search(r'(\d+)[\s]*(?:年|years?|y)', text, re.IGNORECASE)
    if exp_match:
        info['工作经验'] = exp_match.group(1) + "年"
    
    # 期望薪资提取 - 支持中文描述
    salary_match = re.search(r'(?:期望薪资|薪资要求|期望月薪|期望年薪)[\s:：]*([\d\-~～kK万底薪提成薪金工资待遇薪\+＋加]+)', text)
    if not salary_match:
        salary_match = re.search(r'期望薪资[\s]*([\d\-~～kK万底薪提成薪金工资待遇薪\+＋加]+)', text)
    if salary_match:
        info['期望薪资'] = salary_match.group(1).strip()
    
    # 求职岗位提取 - 增强版
    position_patterns = [
        r'(?:求职意向|应聘职位|申请职位|期望职位|目标岗位|求职岗位)[\s:：]*([\u4e00-\u9fa5A-Za-z0-9（）()、/]+)',
        r'(?:期望工作|意向岗位|岗位意向)[\s:：]*([\u4e00-\u9fa5A-Za-z0-9（）()、/]+)',
        r'(?:申请|应聘|求职|职位)[\s:：]*([\u4e00-\u9fa5A-Za-z0-9（）()、/]+)',
        r'^[\s]*(?:职位|岗位)[\s:：]*([\u4e00-\u9fa5A-Za-z0-9（）()、/]+)',
    ]
    
    position_found = False
    for pattern in position_patterns:
        match = re.search(pattern, text)
        if match:
            position = match.group(1).strip()
            position = re.sub(r'^[：:\s]+', '', position)
            info['求职岗位'] = position
            position_found = True
            break
    
    # 如果未匹配到，尝试跨行匹配
    if not position_found:
        match = re.search(
            r'(?:求职意向|应聘职位|申请职位|期望职位|目标岗位)[\s:：]*(.*?)(?=\n|$)',
            text,
            re.DOTALL
        )
        if match:
            position = match.group(1).strip()
            position = re.split(r'[\n\r]+', position)[0]
            info['求职岗位'] = position
    
    # 标准化岗位名称
    info['求职岗位'] = normalize_position(info['求职岗位'])
    
    # 联系方式提取 - 匹配各种格式
    # contact_match = re.search(r'(?:电话|手机|联系方式|联系电话)[\s:：]*([\d\-\+\(\) ]{7,15})', text)
    # if not contact_match:
    #     email_match = re.search(r'(?:邮箱|电子邮箱|email)[\s:：]*([\w\.-]+@[\w\.-]+)', text)
    #     if email_match:
    #         info['联系方式'] = email_match.group(1)
            
    contact_match = re.search(r'(电话|手机|联系方式|联系电话)[：:]\s*([\d\-]+)', text)
    if contact_match:
        info['联系方式'] = contact_match.group(2)
    else:
        email_match = re.search(r'邮箱[：:]\s*([\w\.-]+@[\w\.-]+)', text)
        if email_match:
            info['联系方式'] = email_match.group(1)
    
    return info

def calculate_position_similarity(pos1, pos2):
    """计算两个岗位名称的相似度 (0-1)"""
    if not pos1 or not pos2:
        return 0.0
    
    # 完全匹配
    if pos1 == pos2:
        return 1.0
    
    # 包含关系检查
    if pos1 in pos2 or pos2 in pos1:
        return 0.7
    
    # 使用difflib计算序列匹配度
    seq_matcher = difflib.SequenceMatcher(None, pos1, pos2)
    similarity = seq_matcher.ratio()
    
    # 关键词匹配增强
    common_keywords = 0
    keywords = ['开发', '设计', '销售', '管理', '运营', '分析', '测试', '产品', '市场', '客服']
    
    for kw in keywords:
        if kw in pos1 and kw in pos2:
            common_keywords += 1
            # 每有一个共同关键词，增加相似度
            similarity += 0.1
    
    # 限制在0-1范围内
    return min(max(similarity, 0.0), 1.0)

def match_applicant_to_job(applicant_info, job_info):
    """匹配求职者与企业需求 - 关键指标不匹配时大幅降低整体匹配度"""
    match_result = {
        '学历匹配': '未评估',
        '薪资匹配': '未评估',
        '岗位匹配': '未评估',
        '性别匹配': '未评估',
        '工作经验匹配': '未评估',
        '整体匹配度': '未评估',
        '岗位相似度': '0%'
    }
    
    # 学历匹配
    education_levels = {
        '博士': 5, '博士研究生': 5, '博士及以上': 5,
        '硕士': 4, '硕士研究生': 4, '硕士及以上': 4,
        '本科': 3, '学士': 3, '大学': 3, '本科及以上': 3,
        '大专': 2, '专科': 2, '大专及以上': 2,
        '高中': 1, '中专': 1, '职高': 1, '高中及以上': 1,
        '初中': 0
    }
    
    def get_education_level(edu_str):
        """从学历字符串中提取核心等级"""
        if edu_str in education_levels:
            return education_levels[edu_str]
        
        if "及以上" in edu_str or "以上" in edu_str:
            core_edu = re.sub(r'[及以]上', '', edu_str)
            if core_edu in education_levels:
                return education_levels[core_edu]
        
        for level in education_levels:
            if level in edu_str and level != "以上":
                return education_levels[level]
        
        return 0
    
    applicant_edu = applicant_info.get('学历', '')
    job_edu = job_info.get('学历要求', '')
    
    if applicant_edu and job_edu:
        applicant_level = get_education_level(applicant_edu)
        job_level = get_education_level(job_edu)
        
        if "及以上" in job_edu or "以上" in job_edu:
            match_result['学历匹配'] = '符合' if applicant_level >= job_level else '不符合'
        else:
            match_result['学历匹配'] = '符合' if applicant_level == job_level else '不符合'
    
    # 薪资匹配
    applicant_salary = applicant_info.get('期望薪资', '')
    job_salary = job_info.get('薪资范围', '')
    
    if applicant_salary and job_salary:
        app_min, app_max = extract_salary_range(applicant_salary)
        job_min, job_max = extract_salary_range(job_salary)
        
        if app_min is not None and app_max is not None and job_min is not None and job_max is not None:
            if app_min >= job_min and app_max <= job_max:
                match_result['薪资匹配'] = '符合'
            elif app_min <= job_max and app_max >= job_min:
                match_result['薪资匹配'] = '部分符合'
            else:
                match_result['薪资匹配'] = '不符合'
        else:
            match_result['薪资匹配'] = '无法评估'
    
    # 岗位匹配
    applicant_position = applicant_info.get('求职岗位', '')
    job_position = job_info.get('招聘岗位', '')
    
    # 计算岗位相似度
    position_similarity = calculate_position_similarity(applicant_position, job_position)
    match_result['岗位相似度'] = f"{position_similarity * 100:.0f}%"
    
    if applicant_position and job_position:
        if position_similarity >= 0.85:
            match_result['岗位匹配'] = '高度符合'
        elif position_similarity >= 0.6:
            match_result['岗位匹配'] = '部分符合'
        else:
            match_result['岗位匹配'] = '不符合'
    else:
        match_result['岗位匹配'] = '未评估'
    
    # 性别匹配
    applicant_gender = applicant_info.get('性别', '')
    job_gender = job_info.get('性别要求', '')
    
    if applicant_gender and job_gender:
        if job_gender == '不限' or job_gender == '无要求' or '不限' in job_gender:
            match_result['性别匹配'] = '符合'
        else:
            match_result['性别匹配'] = '符合' if applicant_gender == job_gender else '不符合'
    
    # 工作经验匹配
    applicant_exp = applicant_info.get('工作经验', '')
    job_exp = job_info.get('工作经验要求', '')
    
    if applicant_exp and job_exp:
        try:
            app_exp_years = int(re.search(r'\d+', applicant_exp).group())
            job_exp_years = int(re.search(r'\d+', job_exp).group())
            match_result['工作经验匹配'] = '符合' if app_exp_years >= job_exp_years else '不符合'
        except:
            match_result['工作经验匹配'] = '无法评估'
    
    # 计算整体匹配度
    critical_fields = ['岗位匹配', '学历匹配']
    normal_fields = ['薪资匹配', '性别匹配', '工作经验匹配']
    
    weight_map = {
        '高度符合': 1.0,
        '符合': 1.0,
        '部分符合': 0.6,
        '不符合': 0.0,
        '无法评估': 0.5,
        '未评估': 0.0
    }
    
    critical_penalty = 0.3
    total_score = 0.0
    max_score = 0.0
    critical_fail = False
    
    for field in critical_fields:
        result = match_result[field]
        if result != '未评估':
            score = weight_map.get(result, 0.0)
            if result == '不符合':
                critical_fail = True
                score *= critical_penalty
            total_score += score * 2.0
            max_score += 2.0
    
    for field in normal_fields:
        result = match_result[field]
        if result != '未评估':
            score = weight_map.get(result, 0.0)
            if critical_fail:
                score *= critical_penalty
            total_score += score
            max_score += 1.0
    
    if max_score > 0:
        match_percentage = int((total_score / max_score) * 100)
        if critical_fail and match_percentage > 50:
            match_percentage = 50
        match_result['整体匹配度'] = f"{match_percentage}%"
    else:
        match_result['整体匹配度'] = '无法计算'
    
    return match_result

def extract_salary_range(salary_str):
    """从薪资字符串中提取数字范围，支持中文描述"""
    salary_str = salary_str.replace(',', '').replace('，', '')
    numbers = []
    
    if '万' in salary_str or 'k' in salary_str.lower():
        num_units = re.findall(r'(\d+\.?\d*)([万kK]?)', salary_str)
        for num, unit in num_units:
            num = float(num)
            if unit == '万':
                num *= 10000
            elif unit.lower() == 'k':
                num *= 1000
            numbers.append(num)
    else:
        numbers = [float(num) for num in re.findall(r'\d+\.?\d*', salary_str)]
    
    if numbers:
        if len(numbers) >= 2:
            return min(numbers), max(numbers)
        else:
            return numbers[0], numbers[0]
    
    return None, None

def save_uploaded_file(uploaded_file, file_type):
    """保存上传的文件"""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_ext = uploaded_file.name.split('.')[-1]
    filename = f"{file_type}_{timestamp}.{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    return file_path

def main():
    st.title("就业面试智能体系统")
    st.subheader("求职者与岗位信息匹配平台")
    
    # 初始化session state
    if 'applicant_info' not in st.session_state:
        st.session_state.applicant_info = {}
    if 'job_info' not in st.session_state:
        st.session_state.job_info = {}
    if 'match_result' not in st.session_state:
        st.session_state.match_result = {}
    
    # 上传功能
    st.sidebar.header("文件上传与岗位选择")
    upload_option = st.sidebar.radio("选择操作", ["岗位信息", "求职者简历"])
    
    if upload_option == "岗位信息":
        st.sidebar.subheader("选择岗位类别")
        
        # 创建岗位选择器 - 默认不选择任何岗位
        job_category = st.sidebar.selectbox(
            "请选择岗位类别",
            ["技术岗", "法务岗", "人事岗", "行政岗", "运营岗", "后勤岗", "销售岗"],
            index=None,  # 不默认选择任何岗位
            placeholder="请选择..."
        )
        
        # 设置岗位信息
        if job_category:
            if job_category in JOB_DATABASE:
                st.session_state.job_info = JOB_DATABASE[job_category]
                st.sidebar.success(f"已选择: {job_category}")
            else:
                st.sidebar.error("未找到该岗位信息")
        else:
            # 清空当前选择的岗位信息
            st.session_state.job_info = {}
            st.sidebar.info("请选择一个岗位类别")
    
    else:  # 求职者简历
        st.sidebar.subheader("上传求职者简历")
        uploaded_file = st.sidebar.file_uploader(
            "上传简历 (doc, docx, pdf, txt)",
            type=["doc", "docx", "pdf", "txt"]
        )
        
        if uploaded_file:
            file_path = save_uploaded_file(uploaded_file, "applicant")
            file_ext = file_path.split('.')[-1].lower()
            
            try:
                if file_ext == "pdf":
                    text = extract_text_from_pdf(file_path)
                elif file_ext in ["doc", "docx"]:
                    text = extract_text_from_docx(file_path)
                else:  # txt
                    text = extract_text_from_txt(file_path)
                
                st.session_state.applicant_info = parse_document(text)
                st.sidebar.success("求职者简历解析成功！")
                
                # 调试信息 - 显示提取的原始文本
                with st.expander("查看提取的原始文本"):
                    st.text_area("原始文本", text, height=300)
                
            except Exception as e:
                st.sidebar.error(f"文件解析错误: {str(e)}")
    
    # 显示解析结果
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("岗位信息")
        if st.session_state.job_info:
            job_df = pd.DataFrame.from_dict(
                st.session_state.job_info, 
                orient='index', 
                columns=['值']
            )
            st.dataframe(job_df)
        else:
            st.info("请选择岗位类别")
    
    with col2:
        st.subheader("求职者信息")
        if st.session_state.applicant_info:
            applicant_df = pd.DataFrame.from_dict(
                st.session_state.applicant_info, 
                orient='index', 
                columns=['值']
            )
            st.dataframe(applicant_df)
        else:
            st.info("请上传求职者简历")
    
    # 匹配按钮
    if st.button("进行匹配分析", use_container_width=True):
        if st.session_state.job_info and st.session_state.applicant_info:
            st.session_state.match_result = match_applicant_to_job(
                st.session_state.applicant_info,
                st.session_state.job_info
            )
            st.success("匹配分析完成！")
        else:
            st.warning("请先选择岗位和上传求职者简历")
    
    # 显示匹配结果
    if st.session_state.match_result:
        st.subheader("匹配分析结果")
        
        # 创建结果数据框，排除岗位相似度（将在后面单独显示）
        display_result = {k: v for k, v in st.session_state.match_result.items() if k != '岗位相似度'}
        match_df = pd.DataFrame.from_dict(
            display_result, 
            orient='index', 
            columns=['结果']
        )
        st.dataframe(match_df)
        
        # 显示岗位相似度详情
        applicant_position = st.session_state.applicant_info.get('求职岗位', '无')
        job_position = st.session_state.job_info.get('招聘岗位', '无')
        similarity = st.session_state.match_result.get('岗位相似度', '0%')
        
        st.write(f"**岗位匹配详情**:")
        st.write(f"- 求职者岗位: `{applicant_position}`")
        st.write(f"- 企业岗位: `{job_position}`")
        st.write(f"- 岗位相似度: `{similarity}`")
        
        # 可视化匹配度
        overall_match = st.session_state.match_result.get('整体匹配度', '0%')
        
        # 只有当匹配度是百分比时才显示进度条
        if '%' in overall_match:
            try:
                match_percentage = int(overall_match.strip('%'))
                st.metric("整体匹配度", overall_match)
                st.progress(match_percentage / 100)
                
                # 关键指标检查
                critical_fail = False
                critical_fields = ['岗位匹配', '学历匹配']
                for field in critical_fields:
                    result = st.session_state.match_result.get(field, '')
                    if '不符合' in result:
                        critical_fail = True
                        st.warning(f"⚠️ 关键指标 '{field}' 不符合要求，匹配度大幅降低")
                
                # 匹配建议
                if match_percentage >= 80:
                    st.success("👍 高度匹配：求职者非常适合该职位")
                elif match_percentage >= 60:
                    st.info("👌 中度匹配：求职者基本符合要求")
                elif match_percentage >= 40:
                    st.warning("⚠️ 低度匹配：存在明显不匹配项")
                else:
                    st.error("❌ 不匹配：求职者与职位要求差距较大")
                    
                # 关键指标不符合时的特殊提示
                if critical_fail and match_percentage > 0:
                    st.error("⛔ 关键指标（岗位/学历）不符合，求职者不符合企业基本要求")
                
            except ValueError:
                st.warning("无法计算匹配度百分比")
        else:
            st.warning(f"匹配度数据异常: {overall_match}")

if __name__ == "__main__":
    main()
